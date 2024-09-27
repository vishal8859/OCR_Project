import cv2
import pytesseract
from pytesseract import Output
import os
import numpy as np
from docx import Document

# Define paths
input_image_path = 'C:/Users/user/Pictures/Screenshots/OCR_Test02.png'  # Replace with your image path
output_image_path = 'C:/Users/user/Downloads/output_image.png'  # Replace with your output image path and filename
output_text_path = 'C:/Users/user/Downloads/output_text.txt'  # Replace with your output text file path and filename
output_docx_path = 'C:/Users/user/Downloads/output_text.docx'  # Replace with your output docx file path and filename

def add_branding(image, text="Code Depot", position=(50, 50), font_scale=1, font_thickness=2,
                 text_color=(255, 255, 255), bg_color=(0, 0, 0)):
    overlay = image.copy()
    alpha = 0.6  # Transparency factor.

    # Get the width and height of the text box
    (text_width, text_height), _ = cv2.getTextSize(text, cv2.FONT_HERSHEY_SIMPLEX, font_scale, font_thickness)
    x, y = position

    # Draw a rectangle and put the text on it
    cv2.rectangle(overlay, (x, y + 10), (x + text_width, y - text_height - 10), bg_color, -1)
    cv2.addWeighted(overlay, alpha, image, 1 - alpha, 0, image)
    cv2.putText(image, text, position, cv2.FONT_HERSHEY_SIMPLEX, font_scale, text_color, font_thickness)

    return image

def save_text_output(data, output_text_path):
    with open(output_text_path, 'w', encoding='utf-8') as file:
        for text in data['text']:
            if text.strip():
                file.write(text.strip() + '\n')
    print(f"OCR text output saved to {output_text_path}")

def save_docx_output(data, output_docx_path):
    doc = Document()
    doc.add_heading('OCR Text Output', 0)

    for text in data['text']:
        if text.strip():
            doc.add_paragraph(text.strip())

    doc.save(output_docx_path)
    print(f"OCR text output saved to {output_docx_path}")

def process_image(input_image_path, output_image_path, output_text_path, output_docx_path):
    # Make sure the output directories exist
    os.makedirs(os.path.dirname(output_image_path), exist_ok=True)
    os.makedirs(os.path.dirname(output_text_path), exist_ok=True)
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)

    try:
        # Read the image
        image = cv2.imread(input_image_path)
        if image is None:
            raise FileNotFoundError(f"Image at {input_image_path} not found.")

        # Perform OCR
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        data = pytesseract.image_to_data(gray, output_type=Output.DICT)

        # Draw bounding boxes around text and update text properties
        n_boxes = len(data['level'])
        for i in range(n_boxes):
            if int(data['conf'][i]) > 60:  # Confidence threshold
                (x, y, w, h) = (data['left'][i], data['top'][i], data['width'][i], data['height'][i])
                image = cv2.rectangle(image, (x, y), (x + w, y + h), (0, 255, 0), 2)
                text = data['text'][i].strip()
                if text:  # Only draw text if there's something to draw
                    font_scale = 0.6  # Adjust font scale if necessary
                    image = cv2.putText(image, text, (x, y - 10), cv2.FONT_HERSHEY_SIMPLEX, font_scale, (255, 255, 255), 6)
                    image = cv2.putText(image, text, (x, y - 10), cv2.FONT_HERSHEY_SIMPLEX, font_scale, (0, 0, 0), 2)

        # Add branding to the image
        image = add_branding(image)

        # Save the annotated image
        cv2.imwrite(output_image_path, image)
        print(f"Annotated image saved to {output_image_path}")

        # Save OCR text output in .txt format
        save_text_output(data, output_text_path)

        # Save OCR text output in .docx format
        save_docx_output(data, output_docx_path)

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    process_image(input_image_path, output_image_path, output_text_path, output_docx_path)
